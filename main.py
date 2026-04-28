from flask import Flask, render_template, request, send_file
from datetime import datetime
import zipfile
import os
import tempfile
import re
import calendar
from docx import Document

app = Flask(__name__)

AGREEMENT_TEMPLATE = "template.docx"
IRR_TEMPLATE = "IRRTEMPLATE.docx"


# ---------- DATE FORMAT ----------
def ordinal(n):
    if 11 <= n % 100 <= 13:
        return f"{n}th"
    return f"{n}{['th','st','nd','rd','th','th','th','th','th','th'][n%10]}"


def format_date_ordinal(date_str):
    dt = datetime.strptime(date_str, "%Y-%m-%d")
    return f"{ordinal(dt.day)} {dt.strftime('%B %Y')}"


# ---------- DOCX REPLACEMENT ----------
def replace_doc(doc, replacements):
    def process_paragraph(p):
        text = "".join(run.text for run in p.runs)
        for k, v in replacements.items():
            text = text.replace(f"{{{{{k}}}}}", str(v))
        if p.runs:
            p.runs[0].text = text
            for i in range(1, len(p.runs)):
                p.runs[i].text = ""

    for p in doc.paragraphs:
        process_paragraph(p)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)

    for s in doc.sections:
        for p in s.header.paragraphs:
            process_paragraph(p)
        for p in s.footer.paragraphs:
            process_paragraph(p)


def generate_doc(template, replacements, filename):
    doc = Document(template)
    replace_doc(doc, replacements)
    temp_dir = tempfile.mkdtemp()
    path = os.path.join(temp_dir, filename)
    doc.save(path)
    return path


# ---------- MONEY ----------
def clean_money(val):
    return float(val.replace("£", "").replace(",", "").strip())


def format_money(val):
    return f"{val:,.2f}"


# ---------- CALCULATIONS ----------
def weekly_rent(monthly):
    return round((monthly * 12) / 52, 2)


def pro_rata(monthly, start_date):
    dt = datetime.strptime(start_date, "%Y-%m-%d")
    days = calendar.monthrange(dt.year, dt.month)[1]
    remaining = days - dt.day + 1
    return round((monthly / days) * remaining, 2)


# ---------- ROUTES ----------
@app.route('/')
def index():
    return render_template("index.html")


@app.route('/generate', methods=['POST'])
def generate():
    try:
        # ---- INPUTS ----
        name = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip()
        mobile = request.form.get('mobile', '').strip()
        kin_name = request.form.get('kin_name', '').strip()
        kin_phone = request.form.get('kin_phone', '').strip()
        kin_email = request.form.get('kin_email', '').strip()
        employer = request.form.get('employer', '').strip()

        start_date = request.form.get('start_date', '').strip()
        end_date = request.form.get('end_date', '').strip()
        rent_input = request.form.get('rent', '').strip()
        deposit_input = request.form.get('deposit', '').strip()
        utilities_input = request.form.get('utilities', '250').strip()

        room = request.form.get('room', '').strip()
        property_addr = request.form.get('property', '').strip()
        ref = request.form.get('ref', '').strip()

        if not all([name, start_date, rent_input, deposit_input, room]):
            return "Missing required fields", 400

        # ---- CLEAN VALUES ----
        rent = clean_money(rent_input)
        deposit = clean_money(deposit_input)
        utilities = clean_money(utilities_input)

        # ---- DATES ----
        today = datetime.now().strftime("%Y-%m-%d")
        date_today = format_date_ordinal(today)
        date_start = format_date_ordinal(start_date)
        date_end = format_date_ordinal(end_date) if end_date else "To be agreed"

        # ---- REF ----
        surname = name.split()[-1].upper()
        if not ref:
            ref = f"{surname}.{property_addr.split()[0].upper()}"

        # ---- CALCULATIONS ----
        weekly = weekly_rent(rent)
        pro = pro_rata(rent, start_date)

        total_due = pro + deposit + utilities
        holding = weekly
        move_in = total_due - holding

        # ---- FORMAT ----
        rent_f = format_money(rent)
        deposit_f = format_money(deposit)
        utilities_f = format_money(utilities)

        weekly_f = format_money(weekly)
        pro_f = format_money(pro)
        total_f = format_money(total_due)
        holding_f = format_money(holding)
        move_in_f = format_money(move_in)

        kin_full = kin_name
        if kin_phone:
            kin_full += f" - {kin_phone}"
        if kin_email:
            kin_full += f" - {kin_email}"

        # ---------- AGREEMENT ----------
        agreement_data = {
            "DATE": date_today,
            "NAME": name,
            "START": date_start,
            "END": date_end,
            "RENT": rent_f,
            "DEPOSIT": deposit_f,
            "ROOM": room.upper(),
            "PROPERTY": property_addr,
            "ADDRESS": property_addr,
            "REF": ref,
        }

        agreement_path = generate_doc(
            AGREEMENT_TEMPLATE,
            agreement_data,
            f"Agreement_{surname}.docx"
        )

        # ---------- IRR ----------
        irr_data = {
            "DATE": date_today,
            "NAME": name,
            "EMAIL": email,
            "MOBILE": mobile,
            "KIN": kin_full,
            "EMPLOYER": employer,
            "ADDRESS": property_addr,
            "ROOM": room.upper(),
            "PW": weekly_f,
            "PCM": rent_f,
            "MONTHS": "12",
            "START": date_start,
            "END": date_end,
            "DEPOSIT": deposit_f,
            "PRO_RATA": pro_f,
            "UTILITIES": utilities_f,
            "TOTAL_DUE": total_f,
            "HOLDING_DEPOSIT": holding_f,
            "MOVE_IN_BALANCE": move_in_f,
        }

        irr_path = generate_doc(
            IRR_TEMPLATE,
            irr_data,
            f"IRR_{surname}.docx"
        )

        # ---------- ZIP ----------
        temp_dir = tempfile.mkdtemp()
        zip_path = os.path.join(temp_dir, f"KPI_{surname}.zip")

        with zipfile.ZipFile(zip_path, 'w') as z:
            z.write(agreement_path, os.path.basename(agreement_path))
            z.write(irr_path, os.path.basename(irr_path))

        return send_file(zip_path, as_attachment=True)

    except Exception as e:
        import traceback
        return f"<pre>{str(e)}\n\n{traceback.format_exc()}</pre>", 500


if __name__ == "__main__":
    app.run(debug=True)
